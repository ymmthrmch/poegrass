import random
from datetime import timedelta
from pathlib import Path

from django.conf import settings
from django.core.exceptions import ValidationError
from django.core.files import File
from django.core.validators import FileExtensionValidator
from django.db import models
from django.db.models import UniqueConstraint
from django.utils import timezone
from docx import Document
from docx.shared import Pt
from docx2pdf import convert

from accounts.models import User
from poegrass.utils import japanese_strftime, make_ruby_whole_sentence


class Event(models.Model):
    title = models.CharField(
        verbose_name="タイトル",
        max_length=63,
        blank=True,
    )
    # generate_files内で使うメソッドadd_titleなどで呼び出すために設定
    default_title = "%Y年%m月%d日（%a）の歌会"
    start_time = models.DateTimeField(
        verbose_name="開始時刻",
        default=timezone.now,
    )
    end_time = models.DateTimeField(
        verbose_name="終了時刻",
        blank=True,
    )
    location = models.CharField(
        verbose_name="場所",
        blank=True,
        max_length=63,
    )
    organizer = models.ForeignKey(
        User,
        verbose_name="司会者",
        on_delete=models.SET_NULL,
        null=True,
        blank=False,
    )
    deadline = models.DateTimeField(
        verbose_name="提出締切",
        blank=True,
    )
    STATUS_CHOICES = [
        ("public", "公開"),
        ("limited", "限定公開"),
        ("private", "非公開"),
    ]
    ann_status = models.CharField(
        verbose_name="告知公開設定",
        max_length=7,
        choices=STATUS_CHOICES,
        default="public",
    )
    rec_status = models.CharField(
        verbose_name="記録公開設定",
        max_length=7,
        choices=STATUS_CHOICES,
        default="private",
    )
    ann_desc = models.TextField(
        verbose_name="告知説明",
        max_length=511,
        blank=True,
        null=True,
    )
    ended = models.BooleanField(
        verbose_name="終了済み",
        default=False,
    )
    rec_desc = models.TextField(
        verbose_name="記録説明",
        max_length=511,
        blank=True,
        null=True,
    )

    def file_path(instance, filename):
        """eisou_doc,eisou_pdfのupload_toを設定"""
        return f"events/{instance.pk}/{filename}"

    eisou_doc = models.FileField(
        blank=True,
        null=True,
        upload_to=file_path,
        validators=[FileExtensionValidator(["docx"])],
    )
    eisou_pdf = models.FileField(
        blank=True,
        null=True,
        upload_to=file_path,
        validators=[FileExtensionValidator(["pdf"])],
    )
    eisou_number = models.PositiveIntegerField(
        verbose_name="詠草一覧版数",
        default=0,
    )

    @property
    def ann_is_public(self):
        return self.ann_status == "public"

    @property
    def ann_is_limited(self):
        return self.ann_status == "limited"

    @property
    def ann_is_private(self):
        return self.ann_status == "private"

    @property
    def rec_is_public(self):
        return self.rec_status == "public"

    @property
    def rec_is_limited(self):
        return self.rec_status == "limited"

    @property
    def rec_is_private(self):
        return self.rec_status == "private"

    @property
    def is_past(self):
        now = timezone.now()
        return now > self.start_time

    def clean(self):
        if self.start_time and self.deadline:
            if self.start_time < self.deadline:
                raise ValidationError(
                    {"deadline": "提出締切は開始時刻よりも後ろにはできません。"}
                )
        if self.start_time and self.end_time:
            if self.end_time < self.start_time:
                raise ValidationError(
                    {"end_time": "終了時刻は開始時刻よりも前にはできません。"}
                )
        if self.end_time and not self.rec_is_private:
            if self.end_time > timezone.now():
                raise ValidationError(
                    {"rec_status": "終了時刻より前には記録は公開できません。"}
                )

    def generate_files(self):
        """
        詠草一覧のdocx,pdfファイルを生成する
        """
        if self.deadline > timezone.now():
            raise ValueError("締切が終了していないため、詠草一覧を生成できません。")
        title = self.title
        organizer = self.organizer.name
        date = japanese_strftime(
            timezone.localtime(self.start_time), "%Y年%m月%d日（%a）"
        )  # 例: イベント日が含まれる場合
        participants = Participant.objects.filter(event=self).exclude(
            user=self.organizer
        )
        participants_and_organizer = Participant.objects.filter(event=self)
        tankas = [
            f"{participant.tanka.content}"
            for participant in participants_and_organizer
            if participant.tanka
        ]

        # ドキュメントの作成
        sample_path = settings.MEDIA_ROOT / "events" / "samples" / "utakai_sample.docx"
        doc = Document(sample_path)

        # タイトルの追加
        doc = self.add_title(doc, title)

        # 参加者の追加
        doc = self.add_info(doc, date, organizer, participants)

        # 詠草の追加
        # 行間の設定．8首以下なら等間隔に，9首以上なら4.0．
        line_spacing = 8.0  # float(32 / len(tankas)) if len(tankas) <= 8 else 4.0
        doc = self.add_tankas(
            doc, tankas, basePoint=11.0, rubyPoint=6.0, line_spacing=line_spacing
        )

        # 保存先のdir作成
        path = settings.MEDIA_ROOT / "events" / Path(str(self.pk))
        path.mkdir(parents=True, exist_ok=True)

        # ドキュメントの保存
        doc_path = path / f"{title}.docx"
        doc.save(doc_path)

        # PDFの生成
        pdf_path = path / f"{title}.pdf"

        # LibreOfficeを使ってpdfへ変換
        # subprocess.run(
        #     ["/Applications/LibreOffice.app/Contents/MacOS/soffice",
        #      "--headless",
        #      "--convert-to", "pdf",
        #      doc_path,
        #      "--outdir", path],
        # )

        # doc2pdfを使ってpdfへ変換
        convert(doc_path, pdf_path)

        # pandocを使ってpdfへ変換
        # pypandoc.convert_file(
        #     doc_path,
        #     'pdf',
        #     outputfile=pdf_path,
        #     extra_args=[
        #         "--pdf-engine=lualatex",
        #         "-V", "documentclass=ltjsarticle",
        #         "-V", "luatexjapresetoptions=hiragino-pron"
        #     ]
        # )

        # 版番に応じてファイル名を決定
        if self.eisou_number == 0:
            file_name = title
        else:
            file_name = f"{title}_ver{self.eisou_number + 1}"

        # eisou_doc,eisou_pdfに保存
        with doc_path.open(mode="rb") as f:
            self.eisou_doc.save(f"{file_name}.docx", File(f), save=True)
        with pdf_path.open(mode="rb") as f:
            self.eisou_pdf.save(f"{file_name}.pdf", File(f), save=True)
        self.eisou_number += 1
        self.save()

        # djangoの動作で名前が変更された時,元のファイルを削除
        if Path(self.eisou_doc.path) != doc_path:
            doc_path.unlink()
        if Path(self.eisou_pdf.path) != pdf_path:
            pdf_path.unlink()

        return doc_path, pdf_path

    def add_title(self, doc, title):
        """タイトルを追加する"""
        head = doc.paragraphs[0]
        if title == japanese_strftime(
            timezone.localtime(self.start_time), self.default_title
        ):
            head_title = "京大短歌歌会　詠草一覧"
        else:
            head_title = title
        title_run = head.add_run(head_title)
        title_run.font.size = Pt(14)
        head.paragraph_format.line_spacing = 1.0

        return doc

    def add_info(self, doc, date, organizer, participants):
        """参加者情報を追加する"""
        info = doc.add_paragraph(f"【日付】{date}\n【司会】{organizer}\n")
        info.runs[0].font.size = Pt(12)
        participant_names = "、".join(
            [participant.name for participant in participants]
        )
        parti_run = info.add_run(f"【参加者】{participant_names}")
        parti_run.font.size = Pt(12)
        info.paragraph_format.space_after = Pt(20.0)

        return doc

    def add_tankas(self, doc, tankas, basePoint=11.0, rubyPoint=6.0, line_spacing=4.0):
        """詠草を追加する"""
        body = doc.add_paragraph()
        random.shuffle(tankas)
        for i, tanka in enumerate(tankas):
            # 最初以外改行する
            if i != 0:
                body.add_run("\n")

            # ルビを振って詠草を追加する
            body = make_ruby_whole_sentence(
                body, f"{i + 1}．{tanka}", basePoint=basePoint, rubyPoint=rubyPoint
            )
        body.paragraph_format.line_spacing = line_spacing

        return doc

    def save(self, *args, **kwargs):
        # イベントのタイトルが未設定の場合、自動生成
        if not self.title and self.start_time:
            date = self.start_time
            self.title = japanese_strftime(date, self.default_title)
        # イベントの締切が未設定の場合、開始時刻の3時間前
        if not self.deadline and self.start_time:
            self.deadline = self.start_time - timedelta(hours=3)
        # イベントの場所が未設定の場合、"未定"
        if not self.location:
            self.location = "未定"
        # イベントの終了時刻が未設定の場合、開始時刻の3時間後
        if not self.end_time:
            self.end_time = self.start_time + timedelta(hours=3)
        # イベントの告知説明が未設定の場合、"特になし"
        if not self.ann_desc:
            self.ann_desc = "特になし"
        # 記録が公開される場合、告知は非公開
        if self.rec_status == "public" or self.rec_status == "limited":
            self.ann_status = "private"
        # eisou_doc,eisou_pdfのupload_toを設定
        if self.eisou_doc and self.pk is None:
            uploaded_file = self.eisou_doc
            self.eisou_doc = None
            super().save(*args, **kwargs)
            self.eisou_doc = uploaded_file
        if self.eisou_pdf and self.pk is None:
            uploaded_file = self.eisou_pdf
            self.eisou_pdf = None
            super().save(*args, **kwargs)
            self.eisou_pdf = uploaded_file

        super().save(*args, **kwargs)

    def __str__(self):
        return self.title


class Tanka(models.Model):
    content = models.TextField(
        verbose_name="詠草",
    )
    author = models.ForeignKey(
        User,
        verbose_name="筆名",
        on_delete=models.CASCADE,
        blank=True,
        null=True,
    )
    guest_author = models.CharField(
        verbose_name="筆名",
        default="",
        max_length=63,
        blank=True,
    )
    STATUS_CHOICES = [
        ("public", "公開"),
        ("limited", "限定公開"),
        ("private", "非公開"),
    ]
    status = models.CharField(
        verbose_name="公開設定", max_length=7, choices=STATUS_CHOICES, default="private"
    )
    created_at = models.DateTimeField(
        verbose_name="作成日時",
        auto_now_add=True,
    )

    def clean(self):
        if not self.author and self.guest_author == "":
            raise ValidationError(
                {"guest_author": "ログインするかゲスト筆名を入力してください。"}
            )
        if self.content == "":
            raise ValidationError({"content": "詠草を入力してください．"})

    def save(self, *args, **kwargs):
        if self.author:
            self.guest_author = ""
        super().save(*args, **kwargs)

    @property
    def is_public(self):
        return self.status == "public"

    @property
    def is_limited(self):
        return self.status == "limited"

    @property
    def is_private(self):
        return self.status == "private"

    def __str__(self):
        return self.content if self.content else ""


class Participant(models.Model):
    user = models.ForeignKey(
        User,
        verbose_name="名前",
        on_delete=models.CASCADE,
        blank=True,
        null=True,
    )
    guest_user = models.CharField(
        verbose_name="名前",
        blank=True,
        max_length=63,
    )
    guest_contact = models.EmailField(
        verbose_name="メールアドレス",
        blank=True,
        null=True,
    )
    event = models.ForeignKey(
        Event,
        verbose_name="参加イベント",
        on_delete=models.CASCADE,
    )
    tanka = models.ForeignKey(
        Tanka,
        verbose_name="詠草",
        null=True,
        on_delete=models.SET_NULL,
    )
    message = models.TextField(
        verbose_name="備考欄",
        default="",
        blank=True,
        max_length=200,
    )
    is_observer = models.BooleanField(
        verbose_name="見学者フラグ",
        default=False,
    )

    @property
    def name(self):
        return self.user.name if self.user else self.guest_user

    # def clean(self):
    #     print("cleanが実行されました．")

    #     print(self.user)
    #     print(self.guest_user)

    #     if not self.user:
    #         if self.guest_user == "":
    #             raise ValidationError(
    #                 {"guest_user": "ログインするか筆名を入力してください。"}
    #             )
    #         if not self.guest_contact:
    #             raise ValidationError(
    #                 {
    #                     "guest_contact": "ログインするかメールアドレスを入力してください。"
    #                 }
    #             )

    #     if self.tanka:
    #         if self.user:
    #             if self.tanka.author != self.user:
    #                 raise ValidationError(
    #                     {"tanka": "詠草の筆名と参加者は一致していなければいけません。"}
    #                 )
    #         elif self.guest_user:
    #             if self.tanka.guest_author != self.guest_user:
    #                 raise ValidationError(
    #                     {"tanka": "詠草の筆名と参加者は一致していなければいけません。"}
    #                 )

    def save(self, *args, **kwargs):
        if self.user:
            self.guest_user = ""
        if not self.user:
            self.is_observer = True
        super().save(*args, **kwargs)

    class Meta:
        constraints = [
            UniqueConstraint(fields=["user", "event"], name="unique_user_event")
        ]


class TankaList(models.Model):
    title = models.CharField(
        verbose_name="タイトル",
        max_length=63,
    )
    owner = models.ForeignKey(
        User,
        verbose_name="作成者",
        on_delete=models.CASCADE,
    )
    tankas = models.ManyToManyField(
        Tanka,
        verbose_name="短歌リスト",
        through="TankaListItem",
    )
    description = models.TextField(
        verbose_name="説明",
        max_length=127,
    )
    is_public = models.BooleanField(
        verbose_name="公開設定",
        default=True,
    )
    created_at = models.DateTimeField(
        verbose_name="作成日時",
        auto_now_add=True,
    )

    def add_tanka(self, tanka):
        if not self.tankalistitem_set.filter(tanka=tanka).exists():
            last_order = self.tankalistitem_set.aggregate(
                max_order=models.Max("order")
            )["max_order"]
            new_order = (last_order or 0) + 1
            TankaListItem.objects.create(tanka_list=self, tanka=tanka, order=new_order)
        else:
            raise ValueError("この短歌は既に追加されています。")

    def __str__(self):
        return self.title


class TankaListItem(models.Model):
    tanka_list = models.ForeignKey(TankaList, on_delete=models.CASCADE)
    tanka = models.ForeignKey(Tanka, on_delete=models.CASCADE)
    order = models.PositiveIntegerField()  # 順序を表すフィールド

    class Meta:
        ordering = ["order"]  # orderフィールドに基づいて順序を並べ替え
        constraints = [
            models.UniqueConstraint(
                fields=["tanka_list", "tanka"], name="unique_tanka_in_list"
            )
        ]

    def __str__(self):
        return f"({self.order}){self.tanka.content[:5]}"
